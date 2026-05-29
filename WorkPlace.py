'''
Функциональное и логическое программирование, лабораторное занятие:

По поводу третьей лабораторной:
1. Реализуйте базовую валидацию данных (например, если не заполнено основное поле «Тема письма», отобразить предупреждение).
2. Реорганизуйте расположение элементов управления (компонентов) таким образом, чтобы пользователю было понятно, какие поля относятся к основному телу письма, а какие — к приложениям, информации об отправителе, информации о получателе и т. д.
3.1 Итоговый интерфейс должен быть удобен для пользователя.
3.2 Либо. Реализовать заполнение информации о письме таким образом, чтобы результат был в виден сразу в соседнем окне предпросмотра, интерактивно.
'''
import tkinter as tk
import tkinter.ttk as ttk
import tkinter.filedialog as tkfd
import tkinter.messagebox as tkms
import docx
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import shutil
import os
from datetime import datetime
import tempfile
import fitz
from PIL import Image, ImageTk
from docx2pdf import convert

placeholders = [['{recipient_post}'], ['{recipient_name}'], ['{letter_body}'], ['{date}'], ['{sender_name}']]
adds = []
def main():
    def create_add(i):
        frame_add = ttk.Frame(notebook, padding="10")
        frame_add.columnconfigure(1, weight=1)
        frame_add.grid(row=0, column=0, sticky='nsew')
        notebook.add(frame_add, text=f'Приложение {i}')
        adds.append([])
        for row_i, text_label in enumerate(["Название приложения: ",
                                            "Кол-во страниц:",
                                            "Текст приложения:"]):
            ttk.Label(frame_add, text=text_label).grid(row=row_i, column=0, pady=[0, 3])
            entry = ttk.Entry(frame_add)
            entry.grid(row=row_i, column=1, sticky='nswe', pady=[0, 3])
            adds[i-1].append(entry)
    def pregenerate(file_mask, file_doc, func):
        if any(arr_ph[1].get().strip() == '' for arr_ph in placeholders):
            action = tkms.askyesno(title='Предупреждение', message='Некоторые из основных элементов пусты\n Желаете продолжить?')
            if action == False: return
        try:
            data = placeholders[3][1].get().strip()
        except:
            action = tkms.askyesno(title='Предупреждение', message='Дата введена неверно.\nНужно вводить в формате dd.mm.yyyy\n Желаете продолжить?')
            if action == False: return
        func(file_mask, file_doc)
    def generate(file_mask, file_doc):
        try:
            if file_mask == '':
                raise Exception('Необходимо выбрать шаблон')
            if file_doc != '': adress = file_doc
            else: adress = os.path.join(os.getcwd(), 'Сгенерированное письмо.docx')
            shutil.copy2(file_mask, adress)
            doc = docx.Document(adress)
            build_doc(doc, adress)
            tkms.showinfo(title='Успех', message=f'Файл успешно сгенерирован по адресу:\n{adress}')
        except Exception as ex:
            print(ex.__repr__())
            tkms.showinfo(title='Ошибка', message=ex.__str__())
    def preview(file_mask, file_doc):
        try:
            if file_mask == '': raise Exception('Необходимо выбрать шаблон')

            temp_docx = tempfile.NamedTemporaryFile( delete=False, suffix='.docx')
            docx_path = temp_docx.name
            temp_docx.close()

            temp_pdf = tempfile.NamedTemporaryFile( delete=False, suffix='.pdf')
            pdf_path = temp_pdf.name
            temp_pdf.close()

            shutil.copy2(file_mask, docx_path)
            doc = docx.Document(docx_path)

            build_doc(doc, docx_path)
            convert(docx_path, pdf_path)

            top = tk.Toplevel(root)
            top.title('Предпросмотр')
            top.geometry('900x700')

            canvas = tk.Canvas(top)
            scrollbar = ttk.Scrollbar( top, orient='vertical', command=canvas.yview)
            scroll_frame = ttk.Frame(canvas)
            scroll_frame.bind( "<Configure>", lambda event: canvas.configure( scrollregion=canvas.bbox("all")))

            canvas.create_window((0, 0), window=scroll_frame, anchor='nw')
            canvas.configure( yscrollcommand=scrollbar.set)
            canvas.pack( side='left', fill='both', expand=True)
            scrollbar.pack( side='right', fill='y')

            pdf = fitz.open(pdf_path)
            images = []

            for page_num in range(len(pdf)):
                page = pdf.load_page(page_num)
                pix = page.get_pixmap(matrix=fitz.Matrix(1, 1))
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                photo = ImageTk.PhotoImage(img)
                images.append(photo)
                label = ttk.Label( scroll_frame, image=photo)
                label.pack(pady=10)

            top.images = images
        except Exception as ex:
            print(ex)
            tkms.showerror(title='Ошибка', message=str(ex))
    def build_doc(doc, adress):
        dict_ph = {}
        fmt_text = None
        for ph in placeholders:
            key = ph[0]
            value = ph[1].get()
            dict_ph[key] = value
        for paragraph in doc.paragraphs:
            text = ''.join(run.text for run in paragraph.runs)
            if any(placeholder in text for placeholder in dict_ph.keys()):
                run = paragraph.runs[0]
                fmt = {
                    'bold': run.bold,
                    'italic': run.italic,
                    'underline': run.underline,
                    'font_name': run.font.name,
                    'font_size': run.font.size,
                    'font_color': run.font.color.rgb
                }
                for placeholder, value in dict_ph.items():
                    text = text.replace(placeholder, value)
                    if placeholder == '{letter_body}': fmt_text = {'font_name': run.font.name, 'font_size': run.font.size}
                paragraph.clear()
                new_run = paragraph.add_run(text)
                if fmt['bold']: new_run.bold = fmt['bold']
                if fmt['italic']: new_run.italic = fmt['italic']
                if fmt['underline']: new_run.underline = fmt['underline']
                if fmt['font_name']: new_run.font.name = fmt['font_name']
                if fmt['font_size']: new_run.font.size = fmt['font_size']
                if fmt['font_color']: new_run.font.color.rgb = fmt['font_color']

                if dict_ph['{letter_body}'] in text and len(adds):
                    run_add = paragraph.add_run('\n\tПриложения: \n')
                    run_add.font.size = fmt_text['font_size']
                    run_add.font.name = fmt_text['font_name']
                    for arr_add in adds:
                        run_add = paragraph.add_run(f' - \"{arr_add[0].get()}\" на {arr_add[1].get()} л.\n')
                        run_add.font.size = fmt_text['font_size']
                        run_add.font.name = fmt_text['font_name']
        for i, add in enumerate(adds): 
            #doc.add_page_break()
            paragraph_title1 = doc.add_paragraph(text='\n\n')
            paragraph_title1.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
            run = paragraph_title1.add_run(f'Приложение {i+1}' if len(adds) > 1 else 'Приложение')
            run.font.size = fmt_text['font_size']
            run.font.name = fmt_text['font_name']

            paragraph_title2 = doc.add_paragraph('\n')
            paragraph_title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph_title2.add_run(add[0].get())
            run.font.size = fmt_text['font_size']
            run.font.name = fmt_text['font_name']

            paragraph_title2 = doc.add_paragraph(text='\n')
            paragraph_title2.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph_title2.add_run(f'\t{add[2].get()}')
            run.font.size = fmt_text['font_size']
            run.font.name = fmt_text['font_name']
        doc.save(adress)
    def validate(new_text, added_text, action_type):
        print(f"Действие: {action_type}")
        print(f"Новый текст: '{new_text}'")
        print(f"Добавленный текст: '{added_text}'")
        print("---\n")
        if action_type == '1' and not added_text in '1234567890.': 
            return False
        return True
    root = tk.Tk()
    root.title("Генерация писем в формате DOCX")

    notebook = ttk.Notebook()
    notebook.grid(row=0, column=0, sticky='nsew')

    frame_root = ttk.Frame(notebook)
    frame_root.grid(row=0, column=0, sticky='nsew')
    notebook.add(frame_root, text='Письмо')

    mask = tk.StringVar()
    doc =  tk.StringVar()
    if os.path.exists(os.path.join(os.getcwd(), 'Шаблон ФиЛП.docx')):
        mask.set(os.path.join(os.getcwd(), 'Шаблон ФиЛП.docx'))

    frame_select = ttk.Frame(frame_root, padding="10")
    frame_select.grid(row=0, column=0)

    ttk.Label(frame_select, text="Шаблон документа:").grid(row=0, column=0)
    mask_entry = ttk.Entry(frame_select, textvariable=mask, width=50)
    mask_entry.grid(row=0, column=1)
    but_mask = ttk.Button(frame_select, text='Выбрать', command=lambda: mask.set(value = tkfd.askopenfilename()))
    but_mask.grid(row=0, column=2)

    ttk.Label(frame_select, text="Исходный документ:").grid(row=1, column=0)
    doc_entry = ttk.Entry(frame_select, textvariable=doc, width=50)
    doc_entry.grid(row=1, column=1)
    but_doc = ttk.Button(frame_select, text='Выбрать', command=lambda: doc.set(value = tkfd.asksaveasfilename()))
    but_doc.grid(row=1, column=2)


    frame_write = ttk.Frame(frame_root, padding="10")
    frame_write.grid(row=1, column=0, sticky='nsew')
    frame_write.columnconfigure(1, weight=1)

    for row_i, text_label in enumerate(["Должность адресата (в род.пад.):",
                                "ФИО адресата (в род.пад.):", 
                                "Тело письма:", 
                                "Дата:",
                                "ФИО адресанта:"]):
        ttk.Label(frame_write, text=text_label).grid(row=row_i, column=0)
        entry = ttk.Entry(frame_write)
        if text_label == 'Дата:':
            entry = ttk.Entry(frame_write, validate='key', validatecommand=(frame_write.register(validate),'%P', '%S', '%d'))
        entry.grid(row=row_i, column=1, sticky='nsew')
        placeholders[row_i].append(entry)

    placeholders[0][1].insert(0, 'Преподавателю')
    placeholders[1][1].insert(0 ,'Хандусенко В.О.')
    placeholders[2][1].insert(0, 'Прошу проверить мою 2 лабораторную работу по курсу "ФиЛП", темой которой было расширение функционала программы. Лабораторная работа выполнена в полном объёме, все пункты задания соблюдены.')
    #placeholders[3][1].insert(0, '25.05.2026')
    placeholders[4][1].insert(0, 'Гераськин А.А.')
    
    ttk.Button(frame_write, text="Создать приложение", command=lambda: create_add(len(notebook.tabs()))).grid(row=5, column=0, columnspan=2, pady=[15, 5], sticky='nsew')

    ttk.Button(frame_write, text="Сгенерировать письмо", command=lambda: pregenerate(mask.get(), doc.get(), generate)).grid(row=6, column=0, pady=[5, 5])
    ttk.Button(frame_write, text="Предпросмотр письма", command=lambda: pregenerate(mask.get(), doc.get(), preview)).grid(row=6, column=1, pady=[5, 5])
    root.mainloop()

main()